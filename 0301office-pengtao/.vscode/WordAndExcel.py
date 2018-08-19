# -*- coding:utf-8 -*-
# office_utils
import os
import re
import openpyxl
from docx import Document

class Word:
    """根据不同的内容，生成word文件，输入内容有：标题，副标题，正文，表格，图片"""
    def __init__(self):
        """初始化Word，建一个Document类的实例"""
        self.doc = Document()

    def heading(self):
        """添加标题"""
        try:
            heading = input('请输入标题')
            level = int(input('标题级别，用数字'))
            self.doc.add_heading(heading, level = level)
        except Exception as e:
            print (e)

    def subtitle_title(self):
        """添加副标题"""
        title = input('请输入副标题')
        self.doc.add_paragraph(title, style = 'Title')
    
    def paragraph(self):
        """添加段落"""
        try:
            paragraph = input('请输入段落内容')
            p = self.doc.add_paragraph(paragraph)
            ask = input('是否在该段落继续添加，是（1）否（0）')
            if ask == '1':
                paragraph = input('请输入段落内容')
                p.add_run(paragraph)
        except Exception as e:
            print (e)
    
    def table(self):
        """添加列表"""
        rows, cols = input('行'), input('列')
        table = self.doc.add_table(rows=rows, cols=cols)

    def picture(self):
        """添加图片"""
        picture = input('输入文件名')
        self.doc.add_picture(picture)

    def save(self):
        """保存文件"""
        name = input('请给文件命名')
        self.doc.save(name + '.docx')
        print('word已生成')

    def create_a_word(self):
        """根据选择内容的输入来创建一个word文件"""
        while True:
            print ('这里有五种输入内容：标题（1），副标题（2），段落（3），表格（4），图片（5）')
            choice = input('请选择你要写的内容')
            if choice == '1':
                self.heading()
            elif choice == '2':
                self.subtitle_title()
            elif choice == '3':
                self.paragraph()
            elif choice == '4':
                self.table()
            else:
                self.picture()
            
            ask = input('继续输入？是（1）否（0）')
            if ask == '0':
                self.save()
                break     

    
class Excel:
    def __init__(self):
        pass

    @staticmethod
    def split_by_year(file_path):
        """根据sheet1的第一列的数据，按年份拆分，放到新的工作表，并以年份命名新的sheet"""
        excel = openpyxl.load_workbook(file_path)
        sh1 = excel.active
        year_list = []

        for row in sh1.rows:
            if re.match(r'(\d{4})-', str(row[0].value)):
                year_list.append(re.match(r'(\d{4})-', str(row[0].value)).group(1))
                year_list = list(set(year_list))
                year_list.sort()
        print('这里有这些年份的比特币价格', year_list)
        a = len(year_list)

        for sh_num in range(a):
            year = input('请输入您要拆分的年份')
            sh2 = excel.create_sheet(title=year)
            index = 2
            sh2['A1'].value = '日期'
            sh2['B1'].value = '价格'

            for row in sh1.rows:
                if re.match(r'(\d{4})-', str(row[0].value)) and re.match(r'(\d{4})-', str(row[0].value)).group(1) == year:
                    sh2['A'+ str(index)].value = row[0].value
                    sh2['B'+ str(index)].value = row[1].value
                    index +=1
# 在拆分后的数据下的最后一行添加平均价格
            end = 'B' + str(index - 1)
            sh2['A' + str(index)].value = 'average price'
            sh2['B' + str(index)].value = '=average(B2' + f':{end})'

            try:
                ask = int(input('如需继续添加年份价格表请按1，退出请按0'))
                if ask == 0 :
                    break
            except Exception as e:
                print(e)

        excel.save('new_btc.xlsx')
        print('已生成新的文件')


def main():
    a = Word()
    a.create_a_word()
    file_path = input('请输入您要操作的EXCEL文件,要求含文件路径')
    Excel.split_by_year(file_path)

if __name__ == '__main__':
    main()
